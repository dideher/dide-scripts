from tkinter import *
from tkinter import filedialog
from tkinter.ttk import *
from tkinter.messagebox import showwarning, showinfo
from openpyxl import *
from openpyxl.utils import get_column_letter
import docx
from docx.shared import Pt
import os
from shutil import copyfile


class GUI():
    def __init__(self):
        self.window = Tk()

        self.window.title("Δημιουργία Βεβαιώσεων Μείωσης Διδακτικού Ωραρίου")
        self.window.resizable(False, False)

        self.db = list()
        self.teachers = list()
        self.teachersData = list()
        self.certificateInfo = list()
        self.newCertificatesCreated = False

        self.openDB()
        self.create_widgets()

    def openDB(self):
        db_file = 'db.xlsx'
        db_bak = 'db.bak'
        if (os.path.exists(db_file)):
            copyfile(db_file, db_bak)
            self.parseXlsxData(db_file, self.db)

    def validateSchool(self):
        typed = self.school.get()

        for item in self.schools:
            if typed == item:
                return True

        return False

    def certificateExists(self, certificate):
        for item in self.db:
            if item == certificate:
                return True

        return False

    def createDocx(self):
        self.docx_filename = ''

        if not self.validateSchool():
            self.school.set('')
            self.cbSchool['values'] = self.schools

        fieldsWithError = ''

        schoolYear = self.schoolYear.get()
        date1 = self.date1.get()
        arProt = self.arProt.get()
        school = self.school.get()
        lastName = self.lastName.get()
        firstName = self.firstName.get()
        lastName_accusative = self.lastName_accusative.get()
        firstName_accusative = self.firstName_accusative.get()
        fathersName = self.fathersName.get()
        specialty = self.specialty.get()
        date2 = self.date2.get()
        hours = self.hours.get()

        if schoolYear == '':
            fieldsWithError += 'Διδακτικό Έτος\n'
        if date1 == '':
            fieldsWithError += 'Ημερομηνία (εγγράφου)\n'
        if arProt == '':
            fieldsWithError += 'Αρ. Πρωτ.\n'
        if school == '':
            fieldsWithError += 'Σχολείο\n'
        if lastName == '':
            fieldsWithError += 'Επώνυμο (ονομαστική)\n'
        if firstName == '':
            fieldsWithError += 'Όνομα (ονομαστική)\n'
        if lastName_accusative == '':
            fieldsWithError += 'Επώνυμο (αιτιατική)\n'
        if firstName_accusative == '':
            fieldsWithError += 'Όνομα (αιτιατική)\n'
        if fathersName == '':
            fieldsWithError += 'Πατρώνυμο\n'
        if specialty == '':
            fieldsWithError += 'Ειδικότητα\n'
        if date2 == '':
            fieldsWithError += 'Ημερομηνία αλλαγής ωραρίου:\n'
        if hours == '':
            fieldsWithError += 'Ώρες'

        if fieldsWithError != '':
            showwarning(title='Συμπλήρωση πεδίων',
                        message=f'Παρακαλώ συμπληρώστε τα παρακάτω πεδία:\n{fieldsWithError}')
            return

        espa = ''
        if self.isESPA.get():
            doc = docx.Document('./template-espa.docx')
            self.docx_filename = f'ΕΣΠΑ-{lastName} {firstName} του {fathersName}.docx'
            espa = 'ΕΣΠΑ'
        else:
            doc = docx.Document('./template.docx')
            self.docx_filename = f'{lastName} {firstName} του {fathersName}.docx'

        certificate = [lastName, firstName, fathersName, specialty,
                       schoolYear, date1, arProt, school,
                       espa, date2, hours]

        if self.certificateExists(certificate):
            showwarning(title='Δημιουργία Βεβαίωσης',
                        message='Η Βεβαίωση με τα συγκεκριμένα στοιχεία έχει ήδη δημιουργηθεί.')
            return

        table = doc.tables[0]

        text_a = f'Ηράκλειο, {date1}'
        text_b = f'Αρ. πρωτ.: {arProt}'
        text_c = "ΠΡΟΣ:"
        text_d = f'{lastName_accusative} {firstName_accusative}'
        text_e = "ΚΟΙΝ:"
        text_f = f'  1)  {school}'
        text_g = "  2)  ΠΥΣΔΕ Ν. ΗΡΑΚΛΕΙΟΥ"

        text_1 = "Σύμφωνα με τις διατάξεις του ν. 4152/2013 υποπ.Θ1, παρ2 και του ν.2413/1996 αρ48, παρ2 και των υπ. αρ. " \
                 "123995/Δ1/20-12-2010, 123948/Δ2/6-9-2013 & 181230/Ε2/11-11-2015 εγκυκλίων του Υπουργείου Παιδείας, " \
                 "Έρευνας και Θρησκευμάτων, σχετικά με το υποχρεωτικό ωράριο διδασκαλίας των εκπαιδευτικών και σύμφωνα " \
                 "με το αρχείο που τηρείται στην Υπηρεσία μας [Ο/Η] "
        text_2 = f'{lastName} {firstName}'
        text_3 = " του "
        text_4 = fathersName
        text_6 = specialty
        text_7 = ", έχει συμπληρώσει την απαιτούμενη εκπαιδευτική προϋπηρεσία και δικαιούται τη μείωση του υποχρεωτικού " \
                 "διδακτικού [ΤΟΥ/ΤΗΣ] ωραρίου από "
        text_8 = f'{date2} σε {hours} ώρες'
        text_9 = f', για το διδακτικό έτος {schoolYear}.'

        if (self.sex.get() == 'Άρρεν'):
            text_1 = text_1.replace('[Ο/Η]', 'ο')
            text_5 = " αναπληρωτής εκπαιδευτικός κλάδου "
            text_7 = text_7.replace('[ΤΟΥ/ΤΗΣ]', 'του')
        else:
            text_1 = text_1.replace('[Ο/Η]', 'η')
            text_5 = " αναπληρώτρια εκπαιδευτικός κλάδου "
            text_7 = text_7.replace('[ΤΟΥ/ΤΗΣ]', 'της')

        cell = table.cell(0, 2)

        cell.text = ''
        paragraph = cell.paragraphs[0]
        self.set_normal_style(paragraph, text_a, space_before=36, space_after=0)
        paragraph = cell.add_paragraph()
        self.set_normal_style(paragraph, text_b, space_after=18)
        paragraph = cell.add_paragraph()
        self.set_normal_style(paragraph, text_c, space_after=0)
        paragraph = cell.add_paragraph()
        self.set_normal_style(paragraph, text_d, space_after=18)
        paragraph = cell.add_paragraph()
        self.set_normal_style(paragraph, text_e, space_after=0)
        paragraph = cell.add_paragraph()
        self.set_normal_style(paragraph, text_f, space_after=0)
        paragraph = cell.add_paragraph()
        self.set_normal_style(paragraph, text_g, space_after=0)

        paragraph = doc.paragraphs[3]
        paragraph.text = ''

        run = paragraph.add_run(text_1)
        run = paragraph.add_run(text_2)
        run.bold = True
        run = paragraph.add_run(text_3)
        run = paragraph.add_run(text_4)
        run.bold = True
        run = paragraph.add_run(text_5)
        run = paragraph.add_run(text_6)
        run.bold = True
        run = paragraph.add_run(text_7)
        run = paragraph.add_run(text_8)
        run.bold = True
        run = paragraph.add_run(text_9)

        self.safe_save(doc, self.docx_filename)

        self.newCertificatesCreated = True
        self.db.append(certificate)
        self.updateCertificateInfo()
        showinfo(title='Δημιουργία Βεβαίωσης', message='Η Βεβαίωση δημιουργήθηκε επιτυχώς.')

    def set_normal_style(self, paragraph, text, space_before=-1, space_after=-1):
        run = paragraph.add_run(text)
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(11)

        if (space_before != -1):
            paragraph.paragraph_format.space_before = Pt(space_before)

        if (space_after != -1):
            paragraph.paragraph_format.space_after = Pt(space_after)

    def clearData(self, fromSelection=False):
        self.btnLastCertificate.configure(state='disabled')

        self.date1.set('')
        self.arProt.set('')
        self.school.set('')
        self.cbSchool['values'] = self.schools
        if not fromSelection:
            self.teacher.set('')
            self.cbTeacher['values'] = self.teachers
        self.sex.set('Άρρεν')
        self.lastName.set('')
        self.firstName.set('')
        self.lastName_accusative.set('')
        self.firstName_accusative.set('')
        self.fathersName.set('')
        self.specialty.set('')
        self.date2.set('')
        self.hours.set('')
        self.isESPA.set(False)

    def safe_save(self, doc, outFile):
        notSaved = True

        outputFile = os.path.join(self.outputDirName.get(), outFile)
        while notSaved:
            try:
                doc.save(outputFile)
            except:
                showwarning(title="Αρχείο σε χρήση...",
                            message="Παρακαλώ κλείστε το αρχείο '{}' ώστε να ολοκληρωθεί η αποθήκευση.".format(
                                outputFile))
            else:
                notSaved = False

    def getOutputDirName(self):
        dName = filedialog.askdirectory(initialdir="./data/", title="Επιλέξτε τον φάκελο που θα αποθηκευτούν τα αρχεία")

        if dName == "":
            return

        self.outputDirName.set(dName)
        self.ntrOutputDirName.configure(state='disabled')
        self.btnOpenOutputDir.configure(state='disabled')
        self.ntrSchoolYear.configure(state='normal')
        self.ntrDate1.configure(state='normal')
        self.ntrArProt.configure(state='normal')
        self.cbSchool.configure(state='normal')

        if (str(self.btnOpenTeachersData['state']) == 'disabled'):
            self.cbTeacher.configure(state='normal')

        self.cbSex.configure(state='readonly')
        self.sex.set("Άρρεν")
        self.ntrLastName.configure(state='normal')
        self.ntrFirstName.configure(state='normal')
        self.ntrLastName_accusative.configure(state='normal')
        self.ntrFirstName_accusative.configure(state='normal')
        self.ntrFathersName.configure(state='normal')
        self.ntrSpecialty.configure(state='normal')
        self.ntrDate2.configure(state='normal')
        self.ntrHours.configure(state='normal')
        self.chkESPA.configure(state='normal')
        self.btnCreateDocx.configure(state='normal')
        self.btnClear.configure(state='normal')

    def schoolFilterList(self, e):
        typed = self.school.get()

        if typed == '':
            self.cbSchool['values'] = self.schools
        else:
            filteredList = list()
            for item in self.schools:
                if typed.upper() in item.upper():
                    filteredList.append(item)

            self.cbSchool['values'] = filteredList

    def teacherFilterList(self, e):
        self.clearData(fromSelection=True)
        typed = self.teacher.get()

        if typed == '':
            self.cbTeacher['values'] = self.teachers
        else:
            filteredList = list()
            for item in self.teachers:
                if typed.upper() in item.upper():
                    filteredList.append(item)

            self.cbTeacher['values'] = filteredList

    def teacherSelected(self, e):
        self.teacherFilterList(e)

        id = int(self.teacher.get().split(": ")[0]) - 1

        self.lastName.set(self.teachersData[id][0])
        self.lastNameFocusOut(e)
        self.firstName.set(self.teachersData[id][1])
        self.firstNameFocusOut(e)
        self.fathersName.set(self.possessiveCase(self.teachersData[id][2]))
        self.specialty.set(self.teachersData[id][3])

        self.updateCertificateInfo()

    def countCertificates(self, id):
        count = 0
        for item in self.db:
            if self.teachersData[id][0] == item[0] and self.teachersData[id][1] == item[1] and \
                    self.teachersData[id][3][:3] == item[3][:3] and self.teachersData[id][3] == item[3]:
                count += 1
                self.certificateInfo = [count, item]

        return count

    def updateCertificateInfo(self):
        id = int(self.teacher.get().split(": ")[0]) - 1

        count = self.countCertificates(id)
        if count != 0:
            self.btnLastCertificate.configure(state='normal')
        else:
            self.btnLastCertificate.configure(state='disabled')

    def showCertificateInfo(self):
        count = self.certificateInfo[0]
        info = self.certificateInfo[1]
        msg = f'Πλήθος Βεβαιώσεων: {count}\n\n'
        msg += f'ΕΠΩΝΥΜΟ:\t\t{info[0]}\n'
        msg += f'ΟΝΟΜΑ:\t\t\t{info[1]}\n'
        msg += f'ΠΑΤΡΩΝΥΜΟ:\t\t{info[2]}\n'
        msg += f'ΚΛΑΔΟΣ:\t\t\t{info[3]}\n'
        msg += f'ΔΙΔΑΚΤΙΚΟ ΕΤΟΣ:\t\t{info[4]}\n'
        msg += f'ΗΜΕΡΟΜΗΝΙΑ:\t\t{info[5]}\n'
        msg += f'ΑΡ. ΠΡΩΤ.:\t\t{info[6]}\n'
        msg += f'ΣΧΟΛΕΙΟ:\t\t{info[7]}\n'
        msg += f'ΠΡΟΓΡΑΜΜΑ:\t\t{info[8]}\n'
        msg += f'ΗΜΕΡ. ΑΛΛΑΓΗΣ ΩΡΑΡΙΟΥ:\t{info[9]}\n'
        msg += f'ΩΡΕΣ:\t\t\t{info[10]}'
        showinfo(title='Στοιχεία Τελευταίας Βεβαίωσης', message=msg)

    def possessiveCase(self, text):
        if text[-2:] == 'ΟΣ':
            return f'{text[:-1]}Υ'
        elif text[-2:] == 'ΑΣ':
            return text[:-1]
        elif text[-2:] == 'ΗΣ':
            return text[:-1]
        elif text[-2:] == 'ΩΝ':
            return f'{text}Α'
        else:
            return text

    def lastNameFocusOut(self, e):
        typed = self.lastName.get()

        self.lastName_accusative.set(typed)

        if self.sex.get() == 'Άρρεν':
            if (len(typed) > 1):
                self.lastName_accusative.set(typed[:-1])

    def firstNameFocusOut(self, e):
        typed = self.firstName.get()

        self.firstName_accusative.set(typed)

        if self.sex.get() == 'Άρρεν':
            if (len(typed) > 1):
                self.firstName_accusative.set(typed[:-1])

    def cbSexSelected(self, e):
        ln = self.lastName.get()
        fn = self.firstName.get()

        self.lastName_accusative.set(ln)
        self.firstName_accusative.set(fn)

        if self.sex.get() == 'Άρρεν':
            if (len(ln) > 1):
                self.lastName_accusative.set(ln[:-1])
            if (len(fn) > 1):
                self.firstName_accusative.set(fn[:-1])

    def createTeachers(self):
        for i, item in enumerate(self.teachersData):
            self.teachers.append(f'{i + 1:04}: {item[0]} {item[1]} του {self.possessiveCase(item[2])}, {item[3]}')

    def getTeachersDataFilename(self):
        fName = filedialog.askopenfilename(initialdir="./data/",
                                           title="Επιλέξτε το αρχείο με τους εκαπιδευτικούς",
                                           filetypes=(("xlsx files", "*.xlsx"), ("all files", "*.*")))

        if fName == "":
            return

        self.teachersDataFilename.set(fName)
        self.ntrTeachersDataFilename.configure(state='disabled')
        self.btnOpenTeachersData.configure(state='disabled')
        self.parseXlsxData(fName, self.teachersData)
        self.createTeachers()
        self.cbTeacher['values'] = self.teachers

        if (str(self.btnCreateDocx['state']) != 'disabled'):
            self.cbTeacher.configure(state='normal')

    def parseXlsxData(self, file, data):
        workbook = load_workbook(filename=file)
        sheet = workbook.active

        for row in sheet.iter_rows(min_row=2):
            entry = list()
            for cell in row:
                if cell.value is None:
                    text = ""
                else:
                    text = str(cell.value).strip()

                entry.append(text)

            data.append(entry)

    def create_widgets(self):
        self.fData = Frame(self.window)

        self.lTeachersData = Label(self.fData, text="Αρχείο εκπαιδευτικών:")
        self.lTeachersData.grid(column=0, row=0, padx=10, pady=5, sticky=E)

        self.teachersDataFilename = StringVar()
        self.ntrTeachersDataFilename = Entry(self.fData, width=128, state='readonly',
                                             textvariable=self.teachersDataFilename)
        self.ntrTeachersDataFilename.grid(column=1, row=0, padx=10, pady=5, sticky=W)

        self.btnOpenTeachersData = Button(self.fData, text="Επιλέξτε αρχείο...", command=self.getTeachersDataFilename)
        self.btnOpenTeachersData.grid(column=2, row=0, padx=10, pady=5)

        self.lOutputDirName = Label(self.fData, text="Φάκελος για αποθήκευση\nτων αρχείων:")
        self.lOutputDirName.grid(column=0, row=1, padx=10, pady=10, sticky=E)

        self.outputDirName = StringVar()
        self.ntrOutputDirName = Entry(self.fData, width=128, state='readonly', textvariable=self.outputDirName)
        self.ntrOutputDirName.grid(column=1, row=1, padx=10, pady=10, sticky=W)

        self.btnOpenOutputDir = Button(self.fData, text="Επιλέξτε φάκελο...", command=self.getOutputDirName)
        self.btnOpenOutputDir.grid(column=2, row=1, padx=10, pady=10)

        self.lSchoolYear = Label(self.fData, text="Διδακτικό Έτος:")
        self.lSchoolYear.grid(column=0, row=2, padx=10, pady=10, sticky=E)

        self.schoolYear = StringVar()
        self.ntrSchoolYear = Entry(self.fData, width=128, state='disabled', textvariable=self.schoolYear)
        self.ntrSchoolYear.grid(column=1, row=2, padx=10, pady=10, sticky=W)

        self.lTeacher = Label(self.fData, text="Εκπαιδευτικός:")
        self.lTeacher.grid(column=0, row=3, padx=10, pady=10, sticky=E)

        self.teacher = StringVar()
        self.cbTeacher = Combobox(self.fData, width=125, textvariable=self.teacher, state='disabled')
        self.cbTeacher.bind("<KeyRelease>", self.teacherFilterList)
        self.cbTeacher.bind("<<ComboboxSelected>>", self.teacherSelected)
        self.cbTeacher.grid(column=1, row=3, padx=10, pady=5, sticky=W)

        self.btnLastCertificate = Button(self.fData, text="Τελευταία Βεβαίωση", command=self.showCertificateInfo,
                                         state='disabled')
        self.btnLastCertificate.grid(column=2, row=3, padx=10, pady=10, sticky=W)

        self.lfDocInfoFrame = LabelFrame(self.fData, text="Στοιχεία Εγγράφου")
        self.lfDocInfoFrame.grid(column=0, row=4, columnspan=3, padx=10, pady=10, sticky=EW)

        self.lDate1 = Label(self.lfDocInfoFrame, text="Ημερομηνία:")
        self.lDate1.grid(column=0, row=0, padx=10, pady=10, sticky=E)

        self.date1 = StringVar()
        self.ntrDate1 = Entry(self.lfDocInfoFrame, width=32, state='disabled', textvariable=self.date1)
        self.ntrDate1.grid(column=1, row=0, padx=10, pady=10, sticky=W)

        self.lArProt = Label(self.lfDocInfoFrame, text="Αρ. Πρωτ.:")
        self.lArProt.grid(column=2, row=0, padx=10, pady=10, sticky=E)

        self.arProt = StringVar()
        self.ntrArProt = Entry(self.lfDocInfoFrame, width=32, state='disabled', textvariable=self.arProt)
        self.ntrArProt.grid(column=3, row=0, padx=10, pady=10, sticky=W)

        self.lSchool = Label(self.lfDocInfoFrame, text="Σχολείο:")
        self.lSchool.grid(column=4, row=0, padx=10, pady=10, sticky=E)

        self.schools = [
            "1ο ΓΥΜΝΑΣΙΟ ΗΡΑΚΛΕΙΟΥ",
            "2ο ΓΥΜΝΑΣΙΟ ΗΡΑΚΛΕΙΟΥ",
            "3ο ΓΥΜΝΑΣΙΟ ΗΡΑΚΛΕΙΟΥ",
            "4ο ΓΥΜΝΑΣΙΟ ΗΡΑΚΛΕΙΟΥ",
            "5ο ΓΥΜΝΑΣΙΟ ΗΡΑΚΛΕΙΟΥ",
            "6ο ΓΥΜΝΑΣΙΟ ΗΡΑΚΛΕΙΟΥ",
            "7ο ΓΥΜΝΑΣΙΟ ΗΡΑΚΛΕΙΟΥ",
            "8ο ΓΥΜΝΑΣΙΟ ΗΡΑΚΛΕΙΟΥ",
            "9ο ΓΥΜΝΑΣΙΟ ΗΡΑΚΛΕΙΟΥ",
            "10ο ΓΥΜΝΑΣΙΟ ΗΡΑΚΛΕΙΟΥ",
            "11ο ΓΥΜΝΑΣΙΟ ΗΡΑΚΛΕΙΟΥ",
            "12ο ΓΥΜΝΑΣΙΟ ΗΡΑΚΛΕΙΟΥ",
            "13ο ΓΥΜΝΑΣΙΟ ΗΡΑΚΛΕΙΟΥ",
            "ΓΥΜΝΑΣΙΟ Ν. ΑΛΙΚΑΡΝΑΣΣΟΥ",
            "ΓΥΜΝΑΣΙΟ ΑΓΙΑΣ ΒΑΡΒΑΡΑΣ",
            "ΓΥΜΝΑΣΙΟ ΑΓΙΟΥ ΜΥΡΩΝΑ",
            "ΓΥΜΝΑΣΙΟ ΑΓΙΩΝ ΔΕΚΑ",
            "ΓΥΜΝΑΣΙΟ ΑΡΚΑΛΟΧΩΡΙΟΥ",
            "ΓΥΜΝΑΣΙΟ ΑΡΧΑΝΩΝ",
            "ΓΥΜΝΑΣΙΟ ΑΣΗΜΙΟΥ",
            "ΓΥΜΝΑΣΙΟ ΒΑΓΙΟΝΙΑΣ",
            "ΓΥΜΝΑΣΙΟ ΒΕΝΕΡΑΤΟΥ",
            "ΓΥΜΝΑΣΙΟ ΒΙΑΝΝΟΥ",
            "ΓΥΜΝΑΣΙΟ ΓΑΖΙΟΥ",
            "ΓΥΜΝΑΣΙΟ ΓΕΡΓΕΡΗΣ",
            "ΓΥΜΝΑΣΙΟ ΓΟΥΒΩΝ",
            "ΓΥΜΝΑΣΙΟ ΕΠΙΣΚΟΠΗΣ",
            "ΓΥΜΝΑΣΙΟ ΖΑΡΟΥ",
            "ΓΥΜΝΑΣΙΟ ΘΡΑΨΑΝΟΥ",
            "ΓΥΜΝΑΣΙΟ ΚΑΣΤΕΛΛΙΟΥ",
            "ΓΥΜΝΑΣΙΟ ΚΡΟΥΣΩΝΑ",
            "ΓΥΜΝΑΣΙΟ Λ. ΧΕΡΣΟΝΗΣΟΥ",
            "ΓΥΜΝΑΣΙΟ ΜΑΛΙΩΝ",
            "ΓΥΜΝΑΣΙΟ ΜΕΛΕΣΩΝ",
            "ΓΥΜΝΑΣΙΟ ΜΟΙΡΩΝ",
            "ΓΥΜΝΑΣΙΟ ΜΟΧΟΥ",
            "ΓΥΜΝΑΣΙΟ ΠΟΜΠΙΑΣ",
            "ΓΥΜΝΑΣΙΟ ΠΡΟΦΗΤΗ ΗΛΙΑ",
            "ΓΥΜΝΑΣΙΟ ΠΥΡΓΟΥ",
            "ΓΥΜΝΑΣΙΟ ΤΕΦΕΛΙΟΥ",
            "ΓΥΜΝΑΣΙΟ ΤΥΛΙΣΟΥ",
            "ΓΥΜΝΑΣΙΟ ΤΥΜΠΑΚΙΟΥ",
            "ΓΥΜΝΑΣΙΟ ΧΑΡΑΚΑ",
            "ΕΙΔΙΚΟ ΓΥΜΝΑΣΙΟ ΗΡΑΚΛΕΙΟΥ",
            "ΕΣΠΕΡΙΝΟ ΓΥΜΝΑΣΙΟ ΗΡΑΚΛΕΙΟΥ",
            "ΕΣΠΕΡΙΝΟ ΓΥΜΝΑΣΙΟ ΤΥΜΠΑΚΙΟΥ",
            "ΚΑΛΛΙΤΕΧΝΙΚΟ ΓΥΜΝΑΣΙΟ",
            "ΜΟΥΣΙΚΟ ΣΧΟΛΕΙΟ - ΓΥΜΝΑΣΙΟ",
            "ΠΡΟΤΥΠΟ ΓΥΜΝΑΣΙΟ ΗΡΑΚΛΕΙΟΥ",
            "ΣΧΟΛΕΙΟ ΕΥΡΩΠΑΪΚΗΣ ΠΑΙΔΕΙΑΣ",
            "1ο ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΗΡΑΚΛΕΙΟΥ",
            "2ο ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΗΡΑΚΛΕΙΟΥ",
            "3ο ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΗΡΑΚΛΕΙΟΥ",
            "4ο ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΗΡΑΚΛΕΙΟΥ",
            "5ο ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΗΡΑΚΛΕΙΟΥ",
            "6ο ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΗΡΑΚΛΕΙΟΥ",
            "7ο ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΗΡΑΚΛΕΙΟΥ",
            "8ο ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΗΡΑΚΛΕΙΟΥ",
            "10ο ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΗΡΑΚΛΕΙΟΥ",
            "11ο ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΗΡΑΚΛΕΙΟΥ",
            "13ο ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΗΡΑΚΛΕΙΟΥ",
            "ΓΕΝΙΚΟ ΛΥΚΕΙΟ Ν.ΑΛΙΚΑΡΝΑΣΣΟΥ",
            "ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΑΓΙΑΣ ΒΑΡΒΑΡΑΣ",
            "ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΑΓΙΟΥ ΜΥΡΩΝΑ",
            "ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΑΓΙΩΝ ΔΕΚΑ",
            "ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΑΡΚΑΛΟΧΩΡΙΟΥ",
            "ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΑΡΧΑΝΩΝ",
            "ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΑΣΗΜΙΟΥ",
            "ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΒΙΑΝΝΟΥ",
            "ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΓΑΖΙΟΥ",
            "ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΓΟΥΒΩΝ",
            "ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΕΠΙΣΚΟΠΗΣ",
            "ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΚΑΣΤΕΛΙΟΥ",
            "ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΚΡΟΥΣΩΝΑ",
            "ΓΕΝΙΚΟ ΛΥΚΕΙΟ Λ. ΧΕΡΣΟΝΗΣΟΥ",
            "ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΜΑΛΙΩΝ",
            "ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΜΕΛΕΣΩΝ",
            "ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΜΟΙΡΩΝ",
            "ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΜΟΧΟΥ",
            "ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΠΟΜΠΙΑΣ",
            "ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΤΥΜΠΑΚΙΟΥ",
            "ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΧΑΡΑΚΑ",
            "ΕΣΠΕΡΙΝΟ ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΗΡΑΚΛΕΙΟΥ",
            "ΠΡΟΤΥΠΟ ΓΕΝΙΚΟ ΛΥΚΕΙΟ ΗΡΑΚΛΕΙΟΥ",
            "1ο ΕΠΑ.Λ. ΗΡΑΚΛΕΙΟΥ",
            "2ο ΕΠΑ.Λ. ΗΡΑΚΛΕΙΟΥ",
            "3ο ΕΠΑ.Λ. ΗΡΑΚΛΕΙΟΥ",
            "4ο ΕΠΑ.Λ. ΗΡΑΚΛΕΙΟΥ (ΕΣΠΕΡΙΝΟ)",
            "5ο ΕΠΑ.Λ. ΗΡΑΚΛΕΙΟΥ",
            "6ο ΕΠΑ.Λ. ΗΡΑΚΛΕΙΟΥ",
            "1ο ΕΠΑ.Λ. ΑΡΚΑΛΟΧΩΡΙΟΥ",
            "1ο ΕΠΑ.Λ. ΜΟΙΡΩΝ",
            "ΕΕΕΕΚ ΗΡΑΚΛΕΙΟΥ",
            "ΕΕΕΕΚ ΤΥΜΠΑΚΙΟΥ",
            "ΕΠΑΛ ΕΙΔΙΚΗΣ ΑΓΩΓΗΣ",
            "1ο Ε.Κ.",
            "2ο Ε.Κ."
        ]

        self.school = StringVar()
        self.cbSchool = Combobox(self.lfDocInfoFrame, width=64, textvariable=self.school, state='disabled')
        self.cbSchool.bind("<KeyRelease>", self.schoolFilterList)
        self.cbSchool.bind("<<ComboboxSelected>>", self.schoolFilterList)
        self.cbSchool['values'] = self.schools
        self.cbSchool.grid(column=5, row=0, padx=10, pady=5)

        self.lfTeacherInfoFrame = LabelFrame(self.fData, text="Στοιχεία Εκπαιδευτικού")
        self.lfTeacherInfoFrame.grid(column=0, row=5, columnspan=3, padx=10, pady=10, sticky=EW)

        self.lSex = Label(self.lfTeacherInfoFrame, text="Φύλο:")
        self.lSex.grid(column=0, row=0, padx=10, pady=10, sticky=E)

        self.sex = StringVar()
        self.cbSex = Combobox(self.lfTeacherInfoFrame, width=57, textvariable=self.sex, state='disabled')
        self.cbSex.bind("<<ComboboxSelected>>", self.cbSexSelected)
        self.cbSex['values'] = ['Άρρεν', 'Θήλυ']
        self.cbSex.grid(column=1, row=0, padx=10, pady=10, sticky=W)

        self.lLastName = Label(self.lfTeacherInfoFrame, text="Επώνυμο (ονομαστική):")
        self.lLastName.grid(column=0, row=1, padx=10, pady=10, sticky=E)

        self.lastName = StringVar()
        self.ntrLastName = Entry(self.lfTeacherInfoFrame, width=60, state='disabled', textvariable=self.lastName)
        self.ntrLastName.bind("<FocusOut>", self.lastNameFocusOut)
        self.ntrLastName.grid(column=1, row=1, padx=10, pady=10, sticky=E)

        self.lFirstName = Label(self.lfTeacherInfoFrame, text="Όνομα (ονομαστική):")
        self.lFirstName.grid(column=0, row=2, padx=10, pady=10, sticky=E)

        self.firstName = StringVar()
        self.ntrFirstName = Entry(self.lfTeacherInfoFrame, width=60, state='disabled', textvariable=self.firstName)
        self.ntrFirstName.bind("<FocusOut>", self.firstNameFocusOut)
        self.ntrFirstName.grid(column=1, row=2, padx=10, pady=10, sticky=E)

        self.lFathersName = Label(self.lfTeacherInfoFrame, text="Πατρώνυμο (γενική):")
        self.lFathersName.grid(column=0, row=3, padx=10, pady=10, sticky=E)

        self.fathersName = StringVar()
        self.ntrFathersName = Entry(self.lfTeacherInfoFrame, width=60, state='disabled', textvariable=self.fathersName)
        self.ntrFathersName.grid(column=1, row=3, padx=10, pady=10, sticky=E)

        self.lSpecialty = Label(self.lfTeacherInfoFrame, text="Ειδικότητα:")
        self.lSpecialty.grid(column=0, row=4, padx=10, pady=10, sticky=E)

        self.specialty = StringVar()
        self.ntrSpecialty = Entry(self.lfTeacherInfoFrame, width=60, state='disabled', textvariable=self.specialty)
        self.ntrSpecialty.grid(column=1, row=4, padx=10, pady=10, sticky=E)

        self.lESPA = Label(self.lfTeacherInfoFrame, text="ΕΣΠΑ")
        self.lESPA.grid(column=2, row=0, padx=10, pady=10, sticky=E)

        self.isESPA = BooleanVar()
        self.chkESPA = Checkbutton(self.lfTeacherInfoFrame, text="(Επιλέξτε για εμφάνιση λογότυπου)",
                                   variable=self.isESPA, state='disabled')
        self.chkESPA.grid(column=3, row=0, padx=10, pady=10, sticky=W)

        self.lLastName_accusative = Label(self.lfTeacherInfoFrame, text="Επώνυμο (αιτιατική):")
        self.lLastName_accusative.grid(column=2, row=1, padx=10, pady=10, sticky=E)

        self.lastName_accusative = StringVar()
        self.ntrLastName_accusative = Entry(self.lfTeacherInfoFrame, width=60, state='disabled',
                                            textvariable=self.lastName_accusative)
        self.ntrLastName_accusative.grid(column=3, row=1, padx=10, pady=10, sticky=E)

        self.lFirstName_accusative = Label(self.lfTeacherInfoFrame, text="Όνομα (αιτιατική):")
        self.lFirstName_accusative.grid(column=2, row=2, padx=10, pady=10, sticky=E)

        self.firstName_accusative = StringVar()
        self.ntrFirstName_accusative = Entry(self.lfTeacherInfoFrame, width=60, state='disabled',
                                             textvariable=self.firstName_accusative)
        self.ntrFirstName_accusative.grid(column=3, row=2, padx=10, pady=10, sticky=E)

        self.lDate2 = Label(self.lfTeacherInfoFrame, text="Ημερομηνία αλλαγής ωραρίου:")
        self.lDate2.grid(column=2, row=3, padx=10, pady=10, sticky=E)

        self.date2 = StringVar()
        self.ntrDate2 = Entry(self.lfTeacherInfoFrame, width=60, state='disabled', textvariable=self.date2)
        self.ntrDate2.grid(column=3, row=3, padx=10, pady=10, sticky=E)

        self.lHours = Label(self.lfTeacherInfoFrame, text="Ώρες:")
        self.lHours.grid(column=2, row=4, padx=10, pady=10, sticky=E)

        self.hours = StringVar()
        self.ntrHours = Entry(self.lfTeacherInfoFrame, width=60, state='disabled', textvariable=self.hours)
        self.ntrHours.grid(column=3, row=4, padx=10, pady=10, sticky=E)

        self.btnCreateDocx = Button(self.fData, text="Δημιουργία Βεβαίωσης", command=self.createDocx, state='disabled')
        self.btnCreateDocx.grid(column=1, row=10, padx=10, pady=10)

        self.btnClear = Button(self.fData, text="Καθαρισμός", command=self.clearData, state='disabled')
        self.btnClear.grid(column=2, row=10, padx=10, pady=10)

        self.fData.pack()

    def saveDB(self):
        if self.newCertificatesCreated:
            header = ['ΕΠΩΝΥΜΟ', 'ΟΝΟΜΑ', 'ΠΑΤΡΩΝΥΜΟ', 'ΚΛΑΔΟΣ',
                      'ΔΙΔΑΚΤΙΚΟ ΕΤΟΣ', 'ΗΜΕΡΟΜΗΝΙΑ', 'ΑΡ. ΠΡΩΤ.', 'ΣΧΟΛΕΙΟ',
                      'ΠΡΟΓΡΑΜΜΑ', 'ΗΜΕΡΟΜΗΝΙΑ ΑΛΛΑΓΗΣ ΩΡΑΡΙΟΥ', 'ΩΡΕΣ']

            wb = Workbook()
            ws = wb.active

            ws.append(header)

            for entry in self.db:
                ws.append(entry)

            column_widths = []
            for row in ws.iter_rows():
                for i, cell in enumerate(row):
                    try:
                        column_widths[i] = max(column_widths[i], len(str(cell.value)))
                    except IndexError:
                        column_widths.append(len(str(cell.value)))

            for i, column_width in enumerate(column_widths):
                ws.column_dimensions[get_column_letter(i + 1)].width = column_width * 1.23

            outputFile = "db.xlsx"

            notSaved = True

            while notSaved:
                try:
                    wb.save(outputFile)
                except:
                    showwarning(title="Αρχείο σε χρήση...",
                                message="Παρακαλώ κλείστε το αρχείο '{}' ώστε να ολοκληρωθεί η αποθήκευση.".format(
                                    outputFile))
                else:
                    notSaved = False

    def onClosing(self):
        self.saveDB()
        self.window.destroy()


gui = GUI()
gui.window.protocol("WM_DELETE_WINDOW", gui.onClosing)
gui.window.mainloop()
